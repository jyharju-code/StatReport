"""
profile.py — parse an EXAMPLE REPORT into a ReportProfile.

This is the statistics-domain analog of EditMyRaw's style.py / build_look_profile:
EditMyRaw reduces a reference *image* to a look profile (Lab mean/std, histograms);
here we reduce a reference *report* to a structure/tone profile (headings, section
order, prose tone, table/figure density).

Supports .pdf, .docx, .md, .markdown, .html/.htm, .txt. The original file path and
mime type are preserved so gemini.py can also hand the whole document to a
multimodal model to read its visual layout — the direct parallel to feeding the
reference image itself.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".md": "text/markdown", ".markdown": "text/markdown",
    ".txt": "text/plain", ".html": "text/html", ".htm": "text/html",
}

EXAMPLE_EXT = set(_MIME)


def is_example(path: str) -> bool:
    return Path(path).suffix.lower() in EXAMPLE_EXT


def _strip_html(text: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    return re.sub(r"\s+\n", "\n", text)


def _read_pdf(path: str) -> tuple:
    try:
        from pypdf import PdfReader
    except Exception:
        return "", []
    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return text, []


def _read_docx(path: str) -> tuple:
    try:
        import docx  # python-docx
    except Exception:
        return "", []
    doc = docx.Document(path)
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style and p.style.name.lower().startswith("heading") and p.text.strip()]
    text = "\n".join(p.text for p in doc.paragraphs)
    return text, headings


def _extract_headings(text: str) -> List[str]:
    headings: List[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^#{1,4}\s+(.*)$", line)            # markdown headings
        if m:
            headings.append(m.group(1).strip())
            continue
        # ALL-CAPS or numbered short lines look like headings
        if len(line) <= 80 and re.match(r"^(\d+(\.\d+)*\.?\s+)?[A-Z0-9][\w \-/&,]+$", line):
            words = line.split()
            if len(words) <= 9 and (line.isupper() or line[0].isupper()):
                headings.append(line)
    # de-dupe preserving order
    seen, out = set(), []
    for h in headings:
        key = h.lower()
        if key not in seen:
            seen.add(key)
            out.append(h)
    return out[:40]


@dataclass
class ReportProfile:
    name: str
    raw_path: str
    mime: str
    headings: List[str] = field(default_factory=list)
    tone_sample: str = ""
    n_tables: int = 0
    n_figures: int = 0
    full_text: str = ""

    @property
    def can_attach(self) -> bool:
        """Whether the original file is worth handing to a multimodal model."""
        return self.mime in ("application/pdf",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    def skeleton_text(self) -> str:
        lines = [f"Example report '{self.name}'."]
        if self.headings:
            lines.append("Section structure (in order):")
            lines += [f"  {i+1}. {h}" for i, h in enumerate(self.headings)]
        lines.append(f"Approx. {self.n_tables} tables and {self.n_figures} figures referenced.")
        if self.tone_sample:
            lines.append("Tone / writing sample:")
            lines.append('"""' + self.tone_sample[:1200] + '"""')
        return "\n".join(lines)


def parse_example(path: str) -> ReportProfile:
    path = str(Path(path).expanduser())
    ext = Path(path).suffix.lower()
    mime = _MIME.get(ext, "text/plain")
    headings: List[str] = []

    if ext == ".pdf":
        text, headings = _read_pdf(path)
    elif ext == ".docx":
        text, headings = _read_docx(path)
    else:
        raw = Path(path).read_text(encoding="utf-8", errors="ignore")
        text = _strip_html(raw) if ext in {".html", ".htm"} else raw

    if not headings:
        headings = _extract_headings(text)

    low = text.lower()
    n_tables = low.count("table") + text.count("|")  # rough
    n_figures = low.count("figure") + low.count("fig.") + low.count("chart")

    # tone sample: first substantial paragraph
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if len(p.strip()) > 80]
    tone_sample = paras[0] if paras else text.strip()[:1200]

    return ReportProfile(
        name=Path(path).stem, raw_path=path, mime=mime, headings=headings,
        tone_sample=tone_sample, n_tables=min(n_tables, 99), n_figures=min(n_figures, 99),
        full_text=text[:20000],
    )
